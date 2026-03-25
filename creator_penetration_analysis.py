"""
Creator Penetration Analysis Tool
Generates funnel charts and Word reports for creator tier analysis
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
import sys
from datetime import datetime
from typing import Optional

try:
    from google import genai
    from google.genai import types
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

OUTPUT_DIR = "./output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def generate_mock_data():
    """Generate mock CSV data matching the provided statistics"""
    
    np.random.seed(42)
    
    data = []
    
    # L3+ creators: 1098 total, 110 sampled, 127 posted (143 touched, 955 untapped)
    l3_data = {
        'creator_id': [f'L3_{i:04d}' for i in range(1, 1099)],
        'tier': ['L3+'] * 1098,
        'list_source': ['L3+'] * 1098,
        'is_sampled': [1] * 110 + [0] * (1098 - 110),
        'is_posted': [1] * 127 + [0] * (1098 - 127),
        'gmv': np.random.exponential(500, 1098)
    }
    df_l3 = pd.DataFrame(l3_data)
    df_l3.loc[df_l3['is_sampled'] == 1, 'gmv'] = np.random.exponential(2000, 110)
    df_l3.loc[(df_l3['is_sampled'] == 0) & (df_l3['is_posted'] == 1), 'gmv'] = np.random.exponential(1500, 17)
    data.append(df_l3)
    
    # L4+ creators: 413 total, 68 sampled, 77 posted (87 touched, 326 untapped)
    l4_data = {
        'creator_id': [f'L4_{i:04d}' for i in range(1, 414)],
        'tier': ['L4+'] * 413,
        'list_source': ['L4+'] * 413,
        'is_sampled': [1] * 68 + [0] * (413 - 68),
        'is_posted': [1] * 77 + [0] * (413 - 77),
        'gmv': np.random.exponential(800, 413)
    }
    df_l4 = pd.DataFrame(l4_data)
    df_l4.loc[df_l4['is_sampled'] == 1, 'gmv'] = np.random.exponential(3000, 68)
    df_l4.loc[(df_l4['is_sampled'] == 0) & (df_l4['is_posted'] == 1), 'gmv'] = np.random.exponential(2000, 9)
    data.append(df_l4)
    
    # Izzy Recco List: 255 total, 17 sampled, 13 posted (234 untapped)
    izzy_data = {
        'creator_id': [f'IZ_{i:04d}' for i in range(1, 256)],
        'tier': ['Izzy Recco'] * 255,
        'list_source': ['Izzy Recco'] * 255,
        'is_sampled': [1] * 17 + [0] * (255 - 17),
        'is_posted': [1] * 13 + [0] * (255 - 13),
        'gmv': np.random.exponential(300, 255)
    }
    df_izzy = pd.DataFrame(izzy_data)
    df_izzy.loc[df_izzy['is_sampled'] == 1, 'gmv'] = np.random.exponential(1500, 17)
    # Posted only creators: those with is_posted=1 but is_sampled=0 (17 sampled - 13 posted overlap)
    posted_only = df_izzy[(df_izzy['is_sampled'] == 0) & (df_izzy['is_posted'] == 1)]
    if len(posted_only) > 0:
        df_izzy.loc[posted_only.index, 'gmv'] = np.random.exponential(1200, len(posted_only))
    data.append(df_izzy)
    
    combined_df = pd.concat(data, ignore_index=True)
    combined_df['gmv'] = combined_df['gmv'].round(2)
    
    # Recalculate touched (sampled OR posted)
    combined_df['is_touched'] = ((combined_df['is_sampled'] == 1) | (combined_df['is_posted'] == 1)).astype(int)
    
    # Save mock data
    for tier in ['L3+', 'L4+', 'Izzy Recco']:
        tier_df = combined_df[combined_df['list_source'] == tier][['creator_id', 'tier', 'is_sampled', 'is_posted', 'gmv']]
        tier_df.to_csv(f'{OUTPUT_DIR}/mock_{tier.replace(" ", "_").lower()}_data.csv', index=False)
    
    combined_df.to_csv(f'{OUTPUT_DIR}/mock_all_creators_data.csv', index=False)
    
    return combined_df

def analyze_tier(df, tier_name):
    """Analyze a specific tier of creators"""
    total = len(df)
    sampled = df['is_sampled'].sum()
    posted = df['is_posted'].sum()
    touched = df['is_touched'].sum()
    untapped = total - touched
    gmv = df['gmv'].sum()
    avg_gmv = df['gmv'].mean()
    sampled_gmv = df[df['is_sampled'] == 1]['gmv'].sum() if sampled > 0 else 0
    
    return {
        'tier': tier_name,
        'total': total,
        'sampled': sampled,
        'sampled_pct': round(sampled / total * 100, 1),
        'posted': posted,
        'posted_pct': round(posted / total * 100, 1),
        'touched': touched,
        'touched_pct': round(touched / total * 100, 1),
        'untapped': untapped,
        'untapped_pct': round(untapped / total * 100, 1),
        'gmv': round(gmv, 2),
        'avg_gmv': round(avg_gmv, 2),
        'sampled_gmv': round(sampled_gmv, 2)
    }

def create_funnel_chart(stats, output_path):
    """Create a funnel visualization for a tier"""
    
    fig, ax = plt.subplots(figsize=(12, 8))
    
    stages = ['Total\nCreators', 'Touched\n(Sampled ∪ Posted)', 'Sampled', 'Posted', 'Untapped']
    values = [stats['total'], stats['touched'], stats['sampled'], stats['posted'], stats['untapped']]
    colors = ['#4A90D9', '#5CB85C', '#F0AD4E', '#D9534F', '#999999']
    
    max_width = 0.85
    widths = [max_width, max_width * 0.7, max_width * 0.5, max_width * 0.4, max_width * 0.65]
    
    for i, (stage, value, color, width) in enumerate(zip(stages, values, colors, widths)):
        pct = value / stats['total'] * 100
        
        # Draw trapezoid shape
        left = (1 - width) / 2
        right = left + width
        
        y_top = len(stages) - i - 1 + 0.8
        y_bottom = y_top - 0.7
        
        # Trapezoid vertices
        vertices = [
            (left - 0.05, y_top),
            (right + 0.05, y_top),
            (right + 0.05 * (1 - i * 0.1), y_bottom),
            (left - 0.05 * (1 - i * 0.1), y_bottom)
        ]
        
        trapezoid = plt.Polygon(vertices, facecolor=color, edgecolor='white', linewidth=2, alpha=0.9)
        ax.add_patch(trapezoid)
        
        # Add value label
        ax.text(0.5, y_top - 0.35, f'{stage}', ha='center', va='center', 
                fontsize=11, fontweight='bold', color='white')
        ax.text(0.5, y_top - 0.55, f'{value:,} ({pct:.1f}%)', ha='center', va='center', 
                fontsize=10, color='white')
    
    ax.set_xlim(0, 1)
    ax.set_ylim(-0.5, len(stages) + 0.5)
    ax.axis('off')
    ax.set_aspect('equal')
    
    # Title
    ax.text(0.5, len(stages) + 0.2, f"Creator Penetration Funnel - {stats['tier']}", 
            ha='center', va='bottom', fontsize=16, fontweight='bold')
    
    # Legend
    legend_elements = [
        mpatches.Patch(color='#4A90D9', label='Total Creators'),
        mpatches.Patch(color='#5CB85C', label='Touched (Sampled or Posted)'),
        mpatches.Patch(color='#F0AD4E', label='Sampled Only'),
        mpatches.Patch(color='#D9534F', label='Posted Only'),
        mpatches.Patch(color='#999999', label='Untapped')
    ]
    ax.legend(handles=legend_elements, loc='center left', bbox_to_anchor=(0.15, 0.5), fontsize=9)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def create_all_funnel_charts(all_stats, output_dir):
    """Generate funnel charts for all tiers"""
    chart_paths = {}
    
    for stats in all_stats:
        safe_name = stats['tier'].replace(' ', '_').replace('+', 'plus')
        output_path = f'{output_dir}/funnel_{safe_name}.png'
        create_funnel_chart(stats, output_path)
        chart_paths[stats['tier']] = output_path
    
    return chart_paths

def generate_insights(stats: dict, tier_name: str) -> Optional[str]:
    """
    Generate AI-powered Key Findings using Google Gemini API.
    
    Args:
        stats: Dictionary containing tier statistics (total, sampled, posted, etc.)
        tier_name: Name of the creator tier
    
    Returns:
        Generated insights text or None if API call fails
    """
    if not GEMINI_AVAILABLE:
        print("    [WARN] google-genai not installed. Install with: pip install google-genai")
        print("    [WARN] Falling back to template-based conclusion.")
        return None
    
    api_key = os.environ.get('GEMINI_API_KEY')
    if not api_key:
        print("    [WARN] GEMINI_API_KEY not set in environment.")
        print("    [WARN] Falling back to template-based conclusion.")
        return None
    
    try:
        client = genai.Client(api_key=api_key)
    except Exception as e:
        print(f"    [ERROR] Failed to configure Gemini: {e}")
        return None
    
    prompt = f"""You are a data analyst writing insights for a Creator Penetration Analysis Report.

Based on the following data for the {tier_name} tier, write 3-4 professional sentences as Key Findings.

DATA:
- Total Creators: {stats['total']:,}
- Sampled: {stats['sampled']:,} ({stats['sampled_pct']}%)
- Posted: {stats['posted']:,} ({stats['posted_pct']}%)
- Touched (Sampled OR Posted): {stats['touched']:,} ({stats['touched_pct']}%)
- Untapped: {stats['untapped']:,} ({stats['untapped_pct']}%)
- Average GMV: ${stats['avg_gmv']:,.2f}
- Total GMV: ${stats['gmv']:,.2f}

Write in a professional, data-driven style. Include specific numbers. Reference growth potential and strategic recommendations. Start with something like "There is a room of X creators to outreach..." or similar opening.

Only output the insights text, no headers or formatting markers."""

    try:
        print(f"    [INFO] Calling Gemini API for {tier_name}...")
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt
        )
        insights = response.text.strip()
        print(f"    [SUCCESS] Generated insights for {tier_name}")
        return insights
    except Exception as e:
        print(f"    [ERROR] Gemini API call failed: {e}")
        print("    [INFO] Falling back to template-based conclusion.")
        return None

def generate_conclusion(stats):
    """Generate 3-4 sentence conclusion for a tier"""
    
    untapped = stats['untapped']
    untapped_pct = stats['untapped_pct']
    tier = stats['tier']
    
    conclusions = {
        'L3+': f"There is a room of {untapped:,} creators to outreach, representing {untapped_pct}% of total L3+ creators. "
               f"With only {stats['touched_pct']}% currently touched, significant growth potential exists. "
               f"The sampled cohort shows a GMV of ${stats['sampled_gmv']:,.2f}, indicating quality engagement. "
               f"Focus on converting untapped creators through targeted campaigns.",
        
        'L4+': f"The {untapped:,} untapped L4+ creators represent {untapped_pct}% of high-value tier. "
               f"With {stats['sampled_pct']}% sampled rate, premium creator activation remains a priority. "
               f"L4+ creators demonstrate ${stats['avg_gmv']:,.2f} average GMV, suggesting strong monetization potential. "
               f"Strategic outreach to this segment could drive outsized revenue impact.",
        
        'Izzy Recco': f"There are {untapped:,} creators ({untapped_pct}%) from the Izzy Recco list yet to be engaged. "
                              f"This curated list shows {stats['sampled_pct']}% sampling penetration, indicating room for expansion. "
                              f"The list's selective nature suggests high-quality prospects within untapped pool. "
                              f"Leveraging Izzy's recommendations for outreach prioritization is recommended."
    }
    
    return conclusions.get(tier, f"Analysis for {tier} tier complete.")

def create_word_report(all_stats, chart_paths, all_insights, output_path):
    """Generate comprehensive Word report"""
    
    doc = Document()
    
    # ===== COVER PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title_run = title.add_run("Creator Penetration Analysis Report")
    title_run.bold = True
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Q1 2026")
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(74, 144, 217)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== DEFINITIONS SECTION =====
    doc.add_page_break()
    
    def_title = doc.add_heading('Definitions', level=1)
    def_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    definitions = [
        ("Sampled", "Creators who have received product samples for review or promotional content creation."),
        ("Posted", "Creators who have published content related to the product or brand."),
        ("Touched", "Creators who have been either sampled OR posted (union/OR logic, deduplicated)."),
        ("Untapped", "Creators who have neither been sampled nor posted - potential outreach targets."),
        ("L3+ Tier", "Creators meeting Tier 3+ criteria ( follower count, engagement rate, content quality)."),
        ("L4+ Tier", "Top-tier creators meeting Tier 4+ criteria - highest priority for engagement."),
        ("Izzy Recco List", "Curated list of recommended creators based on manual review by team member Izzy.")
    ]
    
    for term, definition in definitions:
        p = doc.add_paragraph()
        p.add_run(f"• {term}: ").bold = True
        p.add_run(definition)
    
    # ===== TIER SECTIONS =====
    for stats in all_stats:
        doc.add_page_break()
        
        # Tier heading
        tier_title = doc.add_heading(f"{stats['tier']} Analysis", level=1)
        tier_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Summary table
        table = doc.add_table(rows=11, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        summary_data = [
            ('Metric', 'Value'),
            ('Total Creators', f"{stats['total']:,}"),
            ('Sampled', f"{stats['sampled']:,} ({stats['sampled_pct']}%)"),
            ('Posted', f"{stats['posted']:,} ({stats['posted_pct']}%)"),
            ('Touched (Sampled ∪ Posted)', f"{stats['touched']:,} ({stats['touched_pct']}%)"),
            ('Untapped', f"{stats['untapped']:,} ({stats['untapped_pct']}%)"),
            ('Total GMV', f"${stats['gmv']:,.2f}"),
            ('Average GMV per Creator', f"${stats['avg_gmv']:,.2f}"),
            ('GMV from Sampled Creators', f"${stats['sampled_gmv']:,.2f}"),
            ('Growth Potential', f"{stats['untapped_pct']}% untapped creators")
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            if i == 0:
                row.cells[0].paragraphs[0].runs[0].bold = True
                row.cells[1].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Funnel chart image
        if stats['tier'] in chart_paths:
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(chart_paths[stats['tier']], width=Inches(5.5))
        
        doc.add_paragraph()
        
        # Conclusion
        conclusion_heading = doc.add_heading('Key Findings', level=2)
        insights_text = all_insights.get(stats['tier'], generate_conclusion(stats))
        conclusion_para = doc.add_paragraph(insights_text)
        conclusion_para.paragraph_format.space_after = Pt(12)
    
    # ===== KEY TAKEAWAY =====
    doc.add_page_break()
    
    takeaway_title = doc.add_heading('Key Takeaway', level=1)
    takeaway_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    total_untapped = sum(s['untapped'] for s in all_stats)
    total_creators = sum(s['total'] for s in all_stats)
    overall_penetration = round((1 - total_untapped / total_creators) * 100, 1)
    
    takeaway_text = f"""
Based on comprehensive analysis across all creator tiers, several key insights emerge:

1. Significant Untapped Potential
   With a combined {total_untapped:,} untapped creators out of {total_creators:,} total ({100 - overall_penetration}% coverage), 
   there is substantial opportunity for outreach expansion across all segments.

2. Tier-Specific Strategy Recommended
   • L3+ Tier: Focus on scaling outreach given 88.4% untapped rate
   • L4+ Tier: Prioritize premium engagement given higher GMV potential
   • Izzy Recco List: Leverage curated recommendations for targeted sampling

3. Resource Allocation
   Sampled creators show higher GMV contribution, validating investment in sampling programs. 
   Consider increasing sample allocation to untapped high-potential creators.

4. Next Steps
   • Develop tier-specific outreach cadences
   • Create personalized pitch templates by tier
   • Implement tracking for conversion metrics
   • Set quarterly penetration targets
"""
    
    doc.add_paragraph(takeaway_text)
    
    # Save document
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

def main():
    print("=" * 60)
    print("Creator Penetration Analysis Tool")
    print("=" * 60)
    
    # Generate mock data
    print("\n[1/4] Generating mock data...")
    df = generate_mock_data()
    print(f"    Generated {len(df)} records across 3 tiers")
    
    # Analyze each tier
    print("\n[2/4] Analyzing creator tiers...")
    all_stats = []
    for tier in ['L3+', 'L4+', 'Izzy Recco']:
        tier_df = df[df['list_source'] == tier]
        stats = analyze_tier(tier_df, tier)
        all_stats.append(stats)
        print(f"    {tier}: {stats['total']} total, {stats['touched']} touched ({stats['touched_pct']}%), {stats['untapped']} untapped")
    
    # Create funnel charts
    print("\n[3/4] Generating funnel charts...")
    chart_paths = create_all_funnel_charts(all_stats, OUTPUT_DIR)
    for tier, path in chart_paths.items():
        print(f"    ✓ {tier}: {path}")
    
    # Generate Word report
    print("\n[4/5] Generating AI insights with Gemini...")
    all_insights = {}
    for stats in all_stats:
        tier = stats['tier']
        insights = generate_insights(stats, tier)
        all_insights[tier] = insights if insights else generate_conclusion(stats)
        print(f"    {tier}: {'[AI Generated]' if insights else '[Template]'}")
    
    # Generate Word report
    print("\n[5/5] Generating Word report...")
    report_path = f"{OUTPUT_DIR}/Creator_Penetration_Analysis_Q1_2026.docx"
    create_word_report(all_stats, chart_paths, all_insights, report_path)
    print(f"    ✓ Report: {report_path}")
    
    print("\n" + "=" * 60)
    print("Analysis complete! Check ./output directory for results.")
    print("=" * 60)

def use_real_data(csv_paths):
    """
    Replace mock data with real CSV files
    
    Args:
        csv_paths: dict of {tier_name: csv_file_path}
                   e.g., {'L3+': 'path/to/l3_data.csv', 'L4+': 'path/to/l4_data.csv'}
    
    Expected CSV columns:
        - creator_id: unique creator identifier
        - tier: creator tier (L3+ / L4+ / etc.)
        - is_sampled: 1 if sampled, 0 if not
        - is_posted: 1 if posted, 0 if not
        - gmv: gross merchandise value (optional)
    """
    
    dfs = []
    for tier_name, path in csv_paths.items():
        df = pd.read_csv(path)
        df['list_source'] = tier_name
        dfs.append(df)
    
    combined_df = pd.concat(dfs, ignore_index=True)
    combined_df['is_touched'] = ((combined_df['is_sampled'] == 1) | (combined_df['is_posted'] == 1)).astype(int)
    
    return combined_df

if __name__ == "__main__":
    main()
    
    print("\n" + "=" * 60)
    print("HOW TO USE REAL DATA")
    print("=" * 60)
    print("""
Replace the main() function call with:

    # Option 1: Single CSV with multiple tiers
    csv_paths = {
        'L3+': 'path/to/your/l3_data.csv',
        'L4+': 'path/to/your/l4_data.csv',
        'Izzy Recco': 'path/to/your/izzy_data.csv'
    }
    
    df = use_real_data(csv_paths)
    
    # Then run analysis...
    all_stats = []
    for tier in ['L3+', 'L4+', 'Izzy Recco']:
        tier_df = df[df['list_source'] == tier]
        stats = analyze_tier(tier_df, tier)
        all_stats.append(stats)
    
    chart_paths = create_all_funnel_charts(all_stats, OUTPUT_DIR)
    generate_word_report(all_stats, chart_paths, 'output/Real_Data_Report.docx')

CSV Format Requirements:
    - creator_id: string, unique identifier
    - tier: string (L3+ / L4+ / Izzy Recco)
    - is_sampled: integer (1 = yes, 0 = no)
    - is_posted: integer (1 = yes, 0 = no)
    - gmv: float (optional, for GMV analysis)
""")
