"""
Flask Web Application for Creator Penetration Analysis
Supports 5-file upload: 3 tier lists + 2 activity files (sampled/posted)
"""

import os
import uuid
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

os.environ['MPLBACKEND'] = 'Agg'

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
from typing import Optional

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__,
            static_folder=os.path.join(BASE_DIR, 'static'),
            template_folder=os.path.join(BASE_DIR, 'templates'))

app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'uploads')
app.config['OUTPUT_FOLDER'] = os.path.join(BASE_DIR, 'output')

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

logger.info(f"App initialized. UPLOAD_FOLDER: {app.config['UPLOAD_FOLDER']}")
logger.info(f"OUTPUT_FOLDER: {app.config['OUTPUT_FOLDER']}")

try:
    from google import genai
    GEMINI_AVAILABLE = True
    logger.info("Gemini API library loaded")
except ImportError:
    GEMINI_AVAILABLE = False
    logger.warning("Gemini API library not available")


def analyze_tier_with_matching(tier_df, sampled_ids, posted_ids, tier_name):
    """Analyze a specific tier by matching against sampled/posted creator_ids"""
    total = len(tier_df)
    
    sampled = tier_df['creator_id'].isin(sampled_ids).sum()
    posted = tier_df['creator_id'].isin(posted_ids).sum()
    touched = tier_df['creator_id'].isin(sampled_ids | posted_ids).sum()
    untapped = total - touched
    
    return {
        'tier': tier_name,
        'total': int(total),
        'sampled': int(sampled),
        'sampled_pct': round(sampled / total * 100, 1),
        'posted': int(posted),
        'posted_pct': round(posted / total * 100, 1),
        'touched': int(touched),
        'touched_pct': round(touched / total * 100, 1),
        'untapped': int(untapped),
        'untapped_pct': round(untapped / total * 100, 1),
        'gmv': 0,
        'avg_gmv': 0,
        'sampled_gmv': 0
    }


def create_funnel_chart(stats, output_path):
    """Create a funnel visualization for a tier"""
    fig, ax = plt.subplots(figsize=(12, 8))
    
    stages = ['Total\nCreators', 'Touched\n(Sampled ∪ Posted)', 'Sampled', 'Posted', 'Untapped']
    values = [stats['total'], stats['touched'], stats['sampled'], stats['posted'], stats['untapped']]
    colors = ['#4A90D9', '#5CB85C', '#F0AD4E', '#D9534F', '#999999']
    
    max_width = 0.85
    widths = [max_width, max_width * 0.7, max_width * 0.5, max_width * 0.4, max_width * 0.65]
    
    for i, (stage, value, color, width) in enumerate(zip(stages, values, colors, widths)):
        pct = value / stats['total'] * 100
        left = (1 - width) / 2
        right = left + width
        y_top = len(stages) - i - 1 + 0.8
        y_bottom = y_top - 0.7
        
        vertices = [
            (left - 0.05, y_top),
            (right + 0.05, y_top),
            (right + 0.05 * (1 - i * 0.1), y_bottom),
            (left - 0.05 * (1 - i * 0.1), y_bottom)
        ]
        
        trapezoid = mpatches.Polygon(vertices, closed=True, facecolor=color, edgecolor='white', linewidth=2, alpha=0.9)
        ax.add_patch(trapezoid)
        ax.text(0.5, y_top - 0.35, f'{stage}', ha='center', va='center', fontsize=11, fontweight='bold', color='white')
        ax.text(0.5, y_top - 0.55, f'{value:,} ({pct:.1f}%)', ha='center', va='center', fontsize=10, color='white')
    
    ax.set_xlim(0, 1)
    ax.set_ylim(-0.5, len(stages) + 0.5)
    ax.axis('off')
    ax.set_aspect('equal')
    ax.text(0.5, len(stages) + 0.2, f"Creator Penetration Funnel - {stats['tier']}", ha='center', va='bottom', fontsize=16, fontweight='bold')
    
    legend_elements = [
        mpatches.Patch(color='#4A90D9', label='Total Creators'),
        mpatches.Patch(color='#5CB85C', label='Touched (Sampled or Posted)'),
        mpatches.Patch(color='#F0AD4E', label='Sampled Only'),
        mpatches.Patch(color='#D9534F', label='Posted Only'),
        mpatches.Patch(color='#999999', label='Untapped')
    ]
    ax.legend(handles=legend_elements, loc='center left', bbox_to_anchor=(0.15, 0.5), fontsize=9)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    logger.info(f"Chart saved: {output_path}")


def generate_insights(stats: dict, tier_name: str, api_key: str) -> Optional[str]:
    """Generate AI-powered Key Findings using Google Gemini API."""
    if not GEMINI_AVAILABLE:
        return None
    
    try:
        client = genai.Client(api_key=api_key)
    except Exception as e:
        logger.error(f"Failed to initialize Gemini client: {e}")
        return None
    
    prompt = f"""You are a data analyst writing insights for a Creator Penetration Analysis Report.

Based on the following data for the {tier_name} tier, write 3-4 professional sentences as Key Findings.

DATA:
- Total Creators: {stats['total']:,}
- Sampled: {stats['sampled']:,} ({stats['sampled_pct']}%)
- Posted: {stats['posted']:,} ({stats['posted_pct']}%)
- Touched (Sampled OR Posted): {stats['touched']:,} ({stats['touched_pct']}%)
- Untapped: {stats['untapped']:,} ({stats['untapped_pct']}%)

Write in a professional, data-driven style. Include specific numbers. Reference growth potential and strategic recommendations. Start with something like "There is a room of X creators to outreach..." or similar opening.

Only output the insights text, no headers or formatting markers."""

    try:
        response = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
        return response.text.strip()
    except Exception as e:
        logger.error(f"Gemini API call failed: {e}")
        return None


def generate_conclusion(stats):
    """Generate template-based conclusion for a tier"""
    untapped = stats['untapped']
    untapped_pct = stats['untapped_pct']
    tier = stats['tier']
    
    conclusions = {
        'L3+': f"There is a room of {untapped:,} creators to outreach, representing {untapped_pct}% of total L3+ creators. "
               f"With only {stats['touched_pct']}% currently touched, significant growth potential exists. "
               f"The sampled cohort includes {stats['sampled']:,} creators ({stats['sampled_pct']}%), while {stats['posted']:,} have posted content. "
               f"Focus on converting untapped creators through targeted campaigns.",
        'L4+': f"The {untapped:,} untapped L4+ creators represent {untapped_pct}% of high-value tier. "
               f"With {stats['sampled_pct']}% sampled rate and {stats['posted_pct']}% posted rate, premium creator activation remains a priority. "
               f"L4+ creators are top priority for brand partnerships. "
               f"Strategic outreach to this segment could drive outsized revenue impact.",
        'Izzy Recco': f"There are {untapped:,} creators ({untapped_pct}%) from the Izzy Recco list yet to be engaged. "
                      f"This curated list shows {stats['sampled_pct']}% sampling penetration, indicating room for expansion. "
                      f"The list's selective nature suggests high-quality prospects within untapped pool. "
                      f"Leveraging Izzy's recommendations for outreach prioritization is recommended."
    }
    
    return conclusions.get(tier, f"Analysis for {tier} tier complete.")


def create_word_report(all_stats, chart_paths, all_insights, output_path):
    """Generate comprehensive Word report"""
    doc = Document()
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title_run = title.add_run("Creator Penetration Analysis Report")
    title_run.bold = True
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Q1 2026")
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(74, 144, 217)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    def_title = doc.add_heading('Definitions', level=1)
    def_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    definitions = [
        ("Sampled", "Creators who have received product samples (matched from brand's sampled data)."),
        ("Posted", "Creators who have published content related to the brand (matched from brand's posted data)."),
        ("Touched", "Creators who are sampled OR posted (union/OR logic, deduplicated)."),
        ("Untapped", "Creators in the tier list who have neither been sampled nor posted."),
        ("L3+ Tier", "Creators meeting Tier 3+ criteria in the brand's creator database."),
        ("L4+ Tier", "Top-tier creators meeting Tier 4+ criteria - highest priority for engagement."),
        ("Izzy Recco List", "Curated list of recommended creators based on manual review.")
    ]
    
    for term, definition in definitions:
        p = doc.add_paragraph()
        p.add_run(f"• {term}: ").bold = True
        p.add_run(definition)
    
    for stats in all_stats:
        doc.add_page_break()
        tier_title = doc.add_heading(f"{stats['tier']} Analysis", level=1)
        tier_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        summary_data = [
            ('Metric', 'Value'),
            ('Total Creators', f"{stats['total']:,}"),
            ('Sampled', f"{stats['sampled']:,} ({stats['sampled_pct']}%)"),
            ('Posted', f"{stats['posted']:,} ({stats['posted_pct']}%)"),
            ('Touched (Sampled ∪ Posted)', f"{stats['touched']:,} ({stats['touched_pct']}%)"),
            ('Untapped', f"{stats['untapped']:,} ({stats['untapped_pct']}%)"),
            ('Growth Potential', f"{stats['untapped_pct']}% untapped creators")
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            if i == 0:
                row.cells[0].paragraphs[0].runs[0].bold = True
                row.cells[1].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        if stats['tier'] in chart_paths:
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(chart_paths[stats['tier']], width=Inches(5.5))
        
        doc.add_paragraph()
        conclusion_heading = doc.add_heading('Key Findings', level=2)
        insights_text = all_insights.get(stats['tier'], generate_conclusion(stats))
        conclusion_para = doc.add_paragraph(insights_text)
        conclusion_para.paragraph_format.space_after = Pt(12)
    
    doc.add_page_break()
    takeaway_title = doc.add_heading('Key Takeaway', level=1)
    takeaway_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    total_untapped = sum(s['untapped'] for s in all_stats)
    total_creators = sum(s['total'] for s in all_stats)
    overall_penetration = round((1 - total_untapped / total_creators) * 100, 1) if total_creators > 0 else 0
    
    takeaway_text = f"""
Based on comprehensive analysis across all creator tiers, several key insights emerge:

1. Significant Untapped Potential
   With a combined {total_untapped:,} untapped creators out of {total_creators:,} total ({100 - overall_penetration}% coverage), 
   there is substantial opportunity for outreach expansion across all segments.

2. Tier-Specific Strategy Recommended
   • L3+ Tier: Focus on scaling outreach given high untapped rate
   • L4+ Tier: Prioritize premium engagement given higher tier status
   • Izzy Recco List: Leverage curated recommendations for targeted sampling

3. Next Steps
   • Develop tier-specific outreach cadences
   • Create personalized pitch templates by tier
   • Implement tracking for conversion metrics
   • Set quarterly penetration targets
"""
    doc.add_paragraph(takeaway_text)
    doc.save(output_path)
    logger.info(f"Report saved: {output_path}")


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/analyze', methods=['POST'])
def analyze():
    session_id = str(uuid.uuid4())
    session_upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
    session_output_dir = os.path.join(app.config['OUTPUT_FOLDER'], session_id)
    os.makedirs(session_upload_dir, exist_ok=True)
    os.makedirs(session_output_dir, exist_ok=True)
    
    logger.info(f"Starting analysis for session: {session_id}")
    
    try:
        api_key = request.form.get('gemini_api_key', '').strip()
        
        sampled_file = request.files.get('sampled_file')
        posted_file = request.files.get('posted_file')
        
        tier_files = {
            'L3+': request.files.get('l3_file'),
            'L4+': request.files.get('l4_file'),
            'Izzy Recco': request.files.get('izzy_file')
        }
        
        if not sampled_file or not posted_file:
            return jsonify({'success': False, 'error': 'Please upload both Sampled and Posted CSV files'})
        
        sampled_filename = secure_filename(sampled_file.filename) if sampled_file.filename else 'sampled.csv'
        posted_filename = secure_filename(posted_file.filename) if posted_file.filename else 'posted.csv'
        
        sampled_filepath = os.path.join(session_upload_dir, sampled_filename)
        posted_filepath = os.path.join(session_upload_dir, posted_filename)
        
        sampled_file.save(sampled_filepath)
        posted_file.save(posted_filepath)
        
        logger.info(f"Saved files: {sampled_filename}, {posted_filename}")
        
        sampled_df = pd.read_csv(sampled_filepath)
        posted_df = pd.read_csv(posted_filepath)
        
        sampled_ids = set(sampled_df['creator_id'].dropna().astype(str))
        posted_ids = set(posted_df['creator_id'].dropna().astype(str))
        
        logger.info(f"Loaded {len(sampled_ids)} sampled, {len(posted_ids)} posted creator IDs")
        
        tier_data = {}
        for tier_name, file in tier_files.items():
            if file and file.filename:
                filename = secure_filename(file.filename)
                if not filename:
                    filename = f"{tier_name.replace(' ', '_')}.csv"
                filepath = os.path.join(session_upload_dir, filename)
                file.save(filepath)
                df = pd.read_csv(filepath)
                tier_data[tier_name] = df
                logger.info(f"Loaded tier {tier_name}: {len(df)} creators")
        
        if not tier_data:
            return jsonify({'success': False, 'error': 'Please upload at least one tier list (L3+, L4+, or Izzy Recco)'})
        
        all_stats = []
        chart_paths = {}
        
        for tier_name, df in tier_data.items():
            stats = analyze_tier_with_matching(df, sampled_ids, posted_ids, tier_name)
            all_stats.append(stats)
            
            safe_name = tier_name.replace(' ', '_').replace('+', 'plus')
            chart_path = os.path.join(session_output_dir, f'funnel_{safe_name}.png')
            create_funnel_chart(stats, chart_path)
            chart_paths[tier_name] = chart_path
        
        all_insights = {}
        for stats in all_stats:
            tier = stats['tier']
            if api_key:
                insights = generate_insights(stats, tier, api_key)
            else:
                insights = None
            all_insights[tier] = insights if insights else generate_conclusion(stats)
        
        report_path = os.path.join(session_output_dir, 'Creator_Penetration_Analysis_Q1_2026.docx')
        create_word_report(all_stats, chart_paths, all_insights, report_path)
        
        logger.info(f"Analysis complete for session: {session_id}")
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'stats': all_stats,
            'report_path': f'/download-report/{session_id}'
        })
    
    except KeyError as e:
        logger.error(f"Missing column in CSV: {e}")
        return jsonify({'success': False, 'error': f'Missing required column: {e}'})
    except Exception as e:
        logger.error(f"Analysis error: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)})


@app.route('/get-chart/<session_id>/<tier_name>')
def get_chart(session_id, tier_name):
    chart_path = os.path.join(app.config['OUTPUT_FOLDER'], session_id, f'funnel_{tier_name}.png')
    logger.info(f"Serving chart: {chart_path}")
    if os.path.exists(chart_path):
        return send_file(chart_path, mimetype='image/png')
    return 'Chart not found', 404


@app.route('/download-report/<session_id>')
def download_report(session_id):
    report_path = os.path.join(app.config['OUTPUT_FOLDER'], session_id, 'Creator_Penetration_Analysis_Q1_2026.docx')
    logger.info(f"Serving report: {report_path}")
    if os.path.exists(report_path):
        return send_file(report_path, as_attachment=True, download_name='Creator_Penetration_Analysis_Q1_2026.docx')
    return 'Report not found', 404


@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Not found'}), 404


@app.errorhandler(500)
def internal_error(e):
    logger.error(f"Internal server error: {e}", exc_info=True)
    return jsonify({'error': 'Internal server error'}), 500


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)
